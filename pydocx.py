"""
python-docx 智能操作docx  格式的word文档
用 Document() 实例化一个文档类 不传参就是新建文档 传参就是给一个docx文档路径 得到包含文档所有内容的实例对象

文档类下面是段落或者表格 paragraphs tables
段落下面是块或者节 runs
表格下面是行或者列 rows cols
行或者列下面是单元格 cells
大家都有 text 属性 内容

查看段落类型是标题还是正文等等 paragraph.style.name->str=="Heading n"
标题有Heading 1 一级标题等 Normal 正文

添加分页符:
document.add_page_break()

设置加粗 斜体 等
run.bold=True
run.italic=True

添加图片 
document.add_picture("path",width=Cm(12),height=Pt(500))

文档中图片的数量 len(document.inline_shapes)

删除图片 
将图片所在段落的块清空 
删除任意对象:对象._element.getparent().remove(对象._element)

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT 表格单元格垂直对齐
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT 段落水平对齐
设置单元格cell居中上对齐
cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.
设置英文字体:run.font.name="Arial"
设置中文字体:run._element.rPr.rFonts.set(qn("w:eastAsia"),"微软雅黑")
全局设置正文字体:
document.styles["Normal"].font.name="Arial"
document.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"),"微软雅黑")
全局设置 标题 n 的字体:
document.styles["Heading n"].font.name="Arial"
document.styles["Heading n"]._element.rPr.rFonts.set(qn("w:eastAsia"),"微软雅黑")

模板:pip install docxtpl
from docxtpl import DocxTemplate

doc=DocxTemplate("path")
doc.render(dict)
doc.save("new_path")

"""
import os
from docx import Document
from docx import  ImagePart
from docx.shared import Pt,Cm,RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docxtpl import DocxTemplate,InlineImage
import jinja2


docx=Document("HL7_1.docx")
tables=docx.tables
print("表格数量：",len(tables))
for table in tables:
    rows=table.rows
    for row in rows:
        for cell in row.cells:
            print(cell.text)
